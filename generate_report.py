from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()

    # Title
    title = doc.add_heading('GridGuard AI: Technical Blueprint & System Architecture', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta Info
    p = doc.add_paragraph()
    p.add_run('System Version: 4.0.2').italic = True
    p.add_run('\nTarget Audience: Cloud Engineers, DevOps specialists, Backend Architects').italic = True
    p.add_run('\nAuthor: NofluxGiven (Shawn Chareka, Project Lead)').italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('---')

    # Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'GridGuard AI is a high-availability, cloud-native grid protection platform. '
        'It identifies, validates, and isolates electricity theft at the distribution pole level '
        'using a sophisticated "Scalpel Isolation" logic. This document details the tech stack '
        'and orchestration methodology that enables sub-300ms responses at continental scale.'
    )

    # Tech Stack Rationale
    doc.add_heading('2. Tech Stack Rationale & Analogies', level=1)
    doc.add_paragraph(
        'For an engineer accustomed to complex cloud systems, it is best to view the tech stack '
        'through the lens of specialized roles:'
    )

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Technology'
    hdr_cells[2].text = 'Role / Analogy'
    hdr_cells[3].text = 'Why it works?'

    items = [
        ('Ingestion Engine', '.NET 9 ASP.NET', 'The Traffic Controller', 'C# with async/await and non-blocking I/O allows thousands of incoming MQTT streams to be coordinated without thread-pool exhaustion.'),
        ('Object Store', 'Huawei OBS', 'The Evidence Locker', 'Storing raw electrical waveforms in a relational DB is inefficient. We use OBS for bulk data, keeping only a "receipt" (URL) in the database.'),
        ('Inference API', 'Python (FastAPI)', 'The Forensic Lab', 'Python is the industry standard for AI. Wrapping the ModelArts inference API in a FastAPI microservice creates a standardized interface.'),
        ('Grid Data Store', 'Huawei TaurusDB', 'The Account Book', 'A high-performance MySQL-compatible database provides the scale needed for multi-terabyte telemetry history.'),
        ('Messaging', 'Huawei IoTDA (MQTT)', 'The Nervous System', 'MQTT provides a lightweight, resilient connection for thousands of edge nodes, even in low-signal rural environments.')
    ]

    for comp, tech, role, why in items:
        row_cells = table.add_row().cells
        row_cells[0].text = comp
        row_cells[1].text = tech
        row_cells[2].text = role
        row_cells[3].text = why

    # Operational Pipeline
    doc.add_heading('3. Operational Pipeline: "The Story of a Signal"', level=1)
    doc.add_paragraph('This represents the lifecycle of a single detection event from hardware to the dashboard.')

    doc.add_heading('3.1 Step 1: Decentralized Detection', level=2)
    doc.add_paragraph(
        'The ESP32-S3 at the pole samples the electrical waveform at high frequency. '
        'It uses local thresholds to calculate the differential between the pole supply and '
        'the sum of metered legal branches.'
    )

    doc.add_heading('3.2 Step 2: The Telemetry Pulse', level=2)
    doc.add_paragraph(
        'When an anomaly exceeds the 5% threshold, the ESP32 publishes a telemetry payload to '
        'the grid/pole/telemetry topic via MQTT. This payload includes current metadata and a '
        'high-resolution waveform snapshot.'
    )

    doc.add_heading('3.3 Step 3: Evidence Archiving', level=2)
    doc.add_paragraph(
        'The .NET backend receives the signal, serializes the waveform into JSON, and uploads it '
        'to Huawei OBS immediately. This ensures we have forensic-level proof of the theft before '
        'further processing.'
    )

    doc.add_heading('3.4 Step 4: AI Validation', level=2)
    doc.add_paragraph(
        'The backend initiates a request to the ModelArts inference engine. The AI performs '
        'CNN-LSTM Pattern Recognition (distinguishing heavy loads from theft signatures) and '
        'XGBoost Probability (checking historical patterns).'
    )

    doc.add_heading('3.5 Step 5: Incident Resolution & Broadcast', level=2)
    doc.add_paragraph(
        'If AI confidence exceeds 90%, the backend updates the TaurusDB event to "Red" and '
        'generates an incident. The SignalR Hub pushes this to all dashboards, updating '
        'the 3D map in milliseconds.'
    )

    # Infrastructure
    doc.add_heading('4. Cloud Infrastructure & DevOps Strategy', level=1)
    
    doc.add_heading('4.1 Continental Scale via O(1) Routing', level=2)
    doc.add_paragraph(
        'The backend implements a Regional Infrastructure Router using a FrozenDictionary. '
        'This allows O(1) lookup of regional Huawei Cloud credentials and endpoints based '
        'on the pole_id prefix (e.g., ZA, DE, US).'
    )

    doc.add_heading('4.2 Containerization & Orchestration', level=2)
    doc.add_paragraph(
        'Services are containerized using Docker and managed with Huawei Cloud CCE, allowing '
        'for horizontal autoscaling based on MQTT ingestion rates.'
    )

    doc.add_heading('4.3 Fail-Safe Architecture', level=2)
    doc.add_paragraph(
        'We follow a "Normally Closed" setup. If cloud connectivity is lost, electricity remains '
        'on for paying customers. Isolation only occurs when theft is confirmed with over 98% confidence.'
    )

    # Deployment
    doc.add_heading('5. Summary Recommendation for Deployment', level=1)
    doc.add_paragraph('1. Monitor MQTT Queue Latency to ensure ingestion is not lagging.')
    doc.add_paragraph('2. Audit OBS Cleanup Policies; move archived waveforms to cold storage after 6 months.')
    doc.add_paragraph('3. Implement TaurusDB Partitioning by Region as deployments grow.')

    doc.add_paragraph('\nCreated for: NofluxGiven | Lead Engineer: Shawn Chareka')

    filename = 'GridGuardAI_Technical_Blueprint.docx'
    doc.save(filename)
    print(f"Document saved as {filename}")

if __name__ == "__main__":
    create_report()
